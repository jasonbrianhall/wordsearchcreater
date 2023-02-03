#!/usr/bin/env python

import random
import string
#import cups
import re
import copy
import docx
import argparse
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def create_word_search(words, size=30):
    """Creates a word search from a list of words and returns the grid as a list of lists."""
    grid = [['-' for _ in range(size)] for _ in range(size)]
    
    # Function to check if word fits in the grid
    def word_fits(word, x, y, direction):
        for i, letter in enumerate(word):
            if x < 0 or x >= size or y < 0 or y >= size:
                return False
            if grid[x][y] != '-' and grid[x][y] != letter:
                return False
            if direction == 'down':
                x += 1
            elif direction == 'right':
                y += 1
            elif direction == 'diagonal':
                x += 1
                y += 1
            elif direction == 'backward_horizontal':
                y -= 1
            elif direction == 'backward_vertical':
                x -= 1
            elif direction == 'backward_diagonal':
                x -= 1
                y -= 1
        return True
    
    for word in words:
        while True:
            direction = random.choice(['down', 'right', 'diagonal', 'backward_horizontal', 'backward_vertical', 'backward_diagonal'])
            x = random.randint(0, size-1)
            y = random.randint(0, size-1)
            if word_fits(word, x, y, direction):
                break
        for i, letter in enumerate(word):
            grid[x][y] = letter
            if direction == 'down':
                x += 1
            elif direction == 'right':
                y += 1
            elif direction == 'diagonal':
                x += 1
                y += 1
            elif direction == 'backward_horizontal':
                y -= 1
            elif direction == 'backward_vertical':
                x -= 1
            elif direction == 'backward_diagonal':
                x -= 1
                y -= 1
    
    return grid

def replace_dashes(grid):
    x=0
    for row in grid:
        y=0
        for character in row:
            if grid[x][y]=="-":
                grid[x][y]=random.choice(string.ascii_letters).lower()
            y+=1
        x+=1

def print_grid(grid):
	"""Prints the grid as a string."""
	for row in grid:
		print(' '.join(row))

def print_words_table(words):
	"""Prints the words in a table format."""
	print('\nWords:')
	for i, word in enumerate(words):
		print(f'{i+1}. {word}')

def create_msword_file(original_grid, grid, title, original_words, output_file):

	# Create a new Word document
	document = docx.Document()

	# Set the font to Lucida Console
	style = document.styles['Normal']
	font = style.font
	#font.name = 'Consolas'

	# Add the header line
	# Create a new paragraph and set its text
	header = document.add_paragraph()
	header_run = header.add_run(title)

	header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
	header_run.bold = True
	header_run.font.size = docx.shared.Pt(14)
	header_run.font.name = "Arial"
	
	# Create the table
	gridlength=len(grid)
	
	paragraph=""
	for i in range(gridlength):
		for j in range(len(grid[i])):
			if j==0:
				paragraph=paragraph + grid[i][j]
			else:
				paragraph=paragraph + " " + grid[i][j]
				
		paragraph=paragraph+"\n"

	# Create a new paragraph and set its text
	puzzle = document.add_paragraph()
	puzzle_run = puzzle.add_run(paragraph)

	# Center the text, make it bold, set the font size to 12, and set the font to Consolas
	puzzle.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
	#puzzle_run.bold = True
	puzzle_run.font.size = docx.shared.Pt(9)
	puzzle_run.font.name = "Consolas"

	'''words=""
	calclength=len(original_words)/3
	print(calclength, original_words)
	table = document.add_table(rows=int(len(original_words)/3+1), cols=3)
	counter=0
	for x in original_words:
		table.cell(int(counter/3),counter%3).text = x
		counter+=1'''
	maxlen=0
	counter=0
	for x in original_words:
		if len(original_words[counter])>maxlen:
			maxlen=len(original_words[counter])
		counter+=1
	calclength=maxlen+5


	grid_string=""
	counter=1
	for x in original_words:
		if not counter%3==0:
			if counter<10:
				grid_string=grid_string+" " + str(counter) + ") " + x + " "*(calclength-len(x))
			else:
				grid_string=grid_string+str(counter) + ") " + x + " "*(calclength-len(x))
		else:
			if counter<10:
				grid_string=grid_string+" " + str(counter) + ") " + x+ "\n"
			else:
				grid_string=grid_string+str(counter) + ") " + x+ "\n"
		counter+=1

	# Create a new paragraph and set its text
	puzzle = document.add_paragraph()
	puzzle_run = puzzle.add_run(grid_string)

	# Center the text, make it bold, set the font size to 12, and set the font to Consolas
	#puzzle.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
	puzzle_run.bold = True
	puzzle_run.font.size = docx.shared.Pt(9)
	puzzle_run.font.name = "Consolas"

	document.add_page_break()

	# Set the font to Lucida Console
	style = document.styles['Normal']
	font = style.font
	#font.name = 'Consolas'

	# Add the header line
	# Create a new paragraph and set its text
	header = document.add_paragraph()
	header_run = header.add_run(title + " (Key)")

	header.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
	header_run.bold = True
	header_run.font.size = docx.shared.Pt(14)
	header_run.font.name = "Arial"
	
	# Create the table
	gridlength=len(original_grid)
	
	paragraph=""
	for i in range(gridlength):
		for j in range(len(original_grid[i])):
			if j==0:
				paragraph=paragraph + original_grid[i][j]
			else:
				paragraph=paragraph + " " + original_grid[i][j]
				
		paragraph=paragraph+"\n"

	# Create a new paragraph and set its text
	puzzle = document.add_paragraph()
	puzzle_run = puzzle.add_run(paragraph)

	# Center the text, make it bold, set the font size to 12, and set the font to Consolas
	puzzle.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
	#puzzle_run.bold = True
	puzzle_run.font.size = docx.shared.Pt(9)
	puzzle_run.font.name = "Consolas"

	'''words=""
	calclength=len(original_words)/3
	print(calclength, original_words)
	table = document.add_table(rows=int(len(original_words)/3+1), cols=3)
	counter=0
	for x in original_words:
		table.cell(int(counter/3),counter%3).text = x
		counter+=1'''
	maxlen=0
	counter=0
	for x in original_words:
		if len(original_words[counter])>maxlen:
			maxlen=len(original_words[counter])
		counter+=1
	calclength=maxlen+5


	grid_string=""
	counter=1
	for x in original_words:
		if not counter%3==0:
			if counter<10:
				grid_string=grid_string+" " + str(counter) + ") " + x + " "*(calclength-len(x))
			else:
				grid_string=grid_string+str(counter) + ") " + x + " "*(calclength-len(x))
		else:
			if counter<10:
				grid_string=grid_string+" " + str(counter) + ") " + x+ "\n"
			else:
				grid_string=grid_string+str(counter) + ") " + x+ "\n"
		counter+=1

	# Create a new paragraph and set its text
	puzzle = document.add_paragraph()
	puzzle_run = puzzle.add_run(grid_string)

	# Center the text, make it bold, set the font size to 12, and set the font to Consolas
	#puzzle.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
	puzzle_run.bold = True
	puzzle_run.font.size = docx.shared.Pt(9)
	puzzle_run.font.name = "Consolas"



	document.save(output_file)

'''def select_printer():
    """Lists available CUPS printers and allows the user to select one."""
    conn = cups.Connection()
    printers = conn.getPrinters()
    print("Available printers:")
    for i, printer in enumerate(printers):
        print("{}: {}".format(i, printer))
    selected_printer = int(input("Select a printer by its number: "))
    return list(printers.keys())[selected_printer]

def print_to_selected_printer(grid, title, words, printer_id):
    """Prints the word search puzzle to the selected printer."""
    conn = cups.Connection()
    grid_string = '\n'.join([' '.join(row) for row in grid])
    grid_string=title + "\n\n" + grid_string + "\n\n"

    maxlen=0
    counter=0
    for x in words:
        if len(words[counter])>maxlen:
            maxlen=len(words[counter])
        counter+=1
    calclength=maxlen+5

    counter=1
    for x in words:
        if not counter%3==0:
            if counter<10:
                grid_string=grid_string+" " + str(counter) + ") " + x + " "*(calclength-len(x))
            else:
                grid_string=grid_string+str(counter) + ") " + x + " "*(calclength-len(x))
        else:
            if counter<10:
                grid_string=grid_string+" " + str(counter) + ") " + x+ "\n"
            else:
                grid_string=grid_string+str(counter) + ") " + x+ "\n"
        counter+=1
    with open("/tmp/print.txt", "w") as f:
        f.write(grid_string)
    conn.printFile(printer_id, "/tmp/print.txt", "Print Job", {})
'''

if __name__ == '__main__':
	parser = argparse.ArgumentParser(description='Process some files.')

	# Add the input file argument
	parser.add_argument('-i', '--input_file', type=str, help='Dictionary File with the first line being the title', required=True)

	# Add the output file argument
	parser.add_argument('-o', '--output_file', type=str, help='MSWord Output File', required=True)

	parser.add_argument('-s', '--size', type=int, help='Puzzle Size (defaults to 30); it also must be greater then zero and if it is too small of a value, it will create an infinite loop', required=False)

	# Parse the command line arguments
	args = parser.parse_args()

	with open(args.input_file, "r") as f:
		words=f.read().split("\n")
	original_words=words.copy()
	original_words.remove(original_words[0])
	original_words.remove("")
	counter=0
	for x in words:
		if counter==0:
			title=words[counter]
			words.remove(words[0])
		else:
			words[counter] = re.sub(r'[^a-z]+', '', x, flags=re.IGNORECASE).lower()
		counter+=1

	counter=0
	for x in words:
		words[counter] = re.sub(r'[^a-z]+', '', x, flags=re.IGNORECASE).lower()
		counter+=1

	words.remove("")
	if not args.size==None and args.size>0:
		grid = create_word_search(words, args.size)
	else:
		grid = create_word_search(words)
	original_grid=[]
	counter=0
	for x in grid:
		original_grid.append(grid[counter].copy())
		counter+=1

	replace_dashes(grid)

	print_grid(grid)
	print_words_table(words)

	#selected_printer = select_printer()

	original_words.sort()
	print("\nCreating MS-Word File " + args.output_file + "\n")
	create_msword_file(original_grid, grid, title, original_words, args.output_file)
	print("Done Creating MS-Word File")

	#print_to_selected_printer(grid, title, original_words, selected_printer)
	#print_to_selected_printer(original_grid, title, original_words, selected_printer)
